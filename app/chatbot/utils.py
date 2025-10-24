from fastapi import UploadFile, HTTPException
import io
from pypdf import PdfReader
from docx import Document as DocxDocument
from app.chatbot.rag import get_vector_store
import tempfile
from app.llm import get_google_response
from fastapi.responses import JSONResponse
from app.chatbot.agents import rag_agent, customer_agent, agency_agent, other_agent
import numpy as np
import joblib
from sklearn.feature_extraction.text import TfidfVectorizer
import logging
from app.config import BEST_MODEL_FILENAME, VECTORIZER_FILENAME  # Removed CLASSES

MAX_BYTES = 200 * 1024 * 1024  # 200 MB limit


def _enforce_size_limit(file: UploadFile, max_bytes: int = MAX_BYTES) -> bytes:
    """
    Read UploadFile stream safely and enforce a maximum size.
    Returns the file content as bytes.
    """
    buf = io.BytesIO()
    total = 0
    chunk_size = 1024 * 1024  # 1 MB

    while True:
        chunk = file.file.read(chunk_size)
        if not chunk:
            break
        total += len(chunk)
        if total > max_bytes:
            raise HTTPException(status_code=413, detail=f"File too large (> {max_bytes//(1024*1024)} MB).")
        buf.write(chunk)

    file.file.seek(0)
    return buf.getvalue()


def _extract_pdf_text(path: str) -> tuple[str, int]:
    reader = PdfReader(path)
    num_pages = len(reader.pages)

    texts = []
    for i in range(num_pages):
        try:
            page = reader.pages[i]
            texts.append(page.extract_text() or "")
        except Exception as e:
            # Continue extracting even if one page fails
            texts.append("")
    return ("\n".join(texts).strip(), num_pages)


def _extract_docx_text(path: str) -> str:
    doc = DocxDocument(path)
    parts = []

    # Paragraphs
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())

    # Tables (flatten cell text)
    for table in doc.tables:
        for row in table.rows:
            cell_texts = []
            for cell in row.cells:
                # Join runs/paragraphs in a cell
                para_texts = [para.text.strip() for para in cell.paragraphs if para.text and para.text.strip()]
                cell_texts.append(" ".join(para_texts).strip())
            row_text = " | ".join([ct for ct in cell_texts if ct])
            if row_text:
                parts.append(row_text)

    return "\n".join(parts).strip()


def _save_to_temp(content: bytes, suffix: str) -> str:
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    with open(tmp.name, "wb") as f:
        f.write(content)
    return tmp.name


class Chatbot:
    @staticmethod
    async def clean_text(text: str) -> str:
        text = text.replace('â€™', "'")
        text = text.replace('â€œ', '"')
        text = text.replace('â€™', '"')
        text = text.replace('Iâ€™m', "I'm")
        return text
    
    @staticmethod
    async def load_assets():
        try:
            vectorizer = joblib.load(VECTORIZER_FILENAME)
            best_model = joblib.load(BEST_MODEL_FILENAME)
            return vectorizer, best_model
        except Exception as e:
            logging.error(f"Load error: {e}")
            raise

    @classmethod
    async def intent_classification(cls, query: str):
        vectorizer, best_model = await cls.load_assets()
        cleaned_query = await cls.clean_text(query)
        vec = vectorizer.transform([cleaned_query])
        pred_idx = best_model.predict(vec)[0]
        classes = {
            0: "ad_booking",
            1: "agency",
            2: "customer",
            3: "other",
            4: "rp",
            5: "refund"
        }
        return classes[pred_idx]
    
    @staticmethod
    async def rag_response(query: str):
        try:
            store = get_vector_store()
            top_k = 5
            source_filter = None
            include_llm_response = True

            # Build filter if source_filter is provided
            filter_dict = None
            if source_filter:
                filter_dict = {"source_type": {"$eq": source_filter}}
            
            results = store.query_by_text(
                query, 
                top_k=top_k,
                filter=filter_dict
            )
            
            
            # Format results
            formatted_results = []
            context_texts = []
            
            for match in results.get('matches', []):
                result_item = {
                    "id": match['id'],
                    "score": match['score'],
                    "metadata": match['metadata']
                }
                formatted_results.append(result_item)
                
                # Collect context for LLM
                if 'text' in match['metadata']:
                    context_texts.append(match['metadata']['text'])
                    print("Added text to context:", match['metadata']['text'][:100] + "...")
                else:
                    print("No 'text' field in metadata for match:", match['id'])
            
            
            
            # Generate LLM response if requested and we have context
            if include_llm_response and context_texts:
                llm_response = await rag_agent(query, context_texts)
            else:
                llm_response = None

            return llm_response
            
            
            
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Query failed: {str(e)}")

    @classmethod
    async def ai_response(cls, query: str, intent: str):
        if intent == "refund":
            message = {
                "label": "would you like to know more about?",
                "options": ["Rajasthan Patrika", "Refund Policies", "Ad Booking"]
            }
            response = "To process a refund, please provide your booking details or contact our support team at support@rajasthanpatrika.com. Refunds are typically processed within 5-7 business days after verification."
            return response, message

        elif intent == "rp":
            response = await cls.rag_response(query)
            message = {
                "label": "would you like to know more about?",
                "options": ["Rajasthan Patrika", "Refund Policies", "Ad Booking"]
            }
            return response, message


        elif intent == "ad_booking":
            message = {
                "label": "For ad booking, please specify the category you are interested in:",
                "options": ["Customer", "Agency"]
            }
            response = "To help you with ad booking, could you please let me know which category you want to book an ad for? (Customer, Agency)"
            return response, message
            
        elif intent == "customer":
            message = {
                "label": "Would you like to know the ad booking flow for any other category like Agency, or want to know about Rajasthan Patrika or refund policy?",
                "options": ["Agency", "Rajasthan Patrika", "Refund Policy"]
            }
            response = await customer_agent(query)
            return response, message
        elif intent == "agency":
            message = {
                "label": "Would you like to know the ad booking flow for any other category like Customer, or want to know about Rajasthan Patrika or refund policy?",
                "options": ["Customer", "Rajasthan Patrika", "Refund Policy"]
            }
            response = await agency_agent(query)
            return response, message
        elif intent == "other":
            message = {
                "label": "Would you like to know the ad booking flow for any other category like Customer, Agency, or want to know about Rajasthan Patrika or refund policy?",
                "options": ["Customer", "Agency", "Rajasthan Patrika", "Refund Policy"]
            }
            response = await other_agent(query)
            return response, message
        else:
            message = {
                "label": "Would you like to know the ad booking flow for any other category like Customer, Agency, or want to know about Rajasthan Patrika or refund policy?",
                "options": ["Customer", "Agency", "Rajasthan Patrika", "Refund Policy"]
            }
            response = await other_agent(query)
            return response, message

